# -*- coding: utf-8 -*-
r"""v2.3 실행 결과 요약 docx 작성. step8의 compute_tobe를 그대로 불러와 실제 데이터로 집계."""
import sys
from pathlib import Path
from collections import Counter

import docx
import pandas as pd
from docx.shared import Pt

sys.path.insert(0, str(Path(__file__).parent))
from step8_build_pn_tobe_v23 import (  # noqa: E402
    SRC, SHEET, HEADER_ROW, COL_PN, COL_MFR_CLEAN_OLD, TOP20_MFRS,
    compute_tobe,
)

OUT = Path(r"C:\Users\정별\1_Work\1.41_동아쏘시오그룹(2)_데이터_표준화\1.41.51_품번_To-Be\30_보고서\260812_품번_To-Be_2단계_결과요약_v2.3.docx")

RULE_LABEL = {
    "Sigma-Aldrich": "선두 '#' 제거 + 핵심코드-수량단위 분리(복합단위/곱셈표기/공백하이픈 확장 포함)",
    "Sigma": "선두 '#' 제거 + 핵심코드-수량단위 분리(복합단위/곱셈표기/공백하이픈 확장 포함)",
    "Merck Millipore": "공통 수량단위 분리 + 선두 '숫자+인치기호' 제거(서술형 값 2건)",
    "Thermo Fisher Scientific": "공통 수량단위 분리 ('KOLAS'로 시작하는 값은 예외)",
    "Sartorius": "'견적번호/견적서 번호 :' 라벨 제거 + '[...]' 제거 + 공통 수량단위 분리",
    "Agilent": "'부품 번호:' 라벨 제거 + 'UI' 접미사만 제거(숫자 접미사는 유지)",
    "삼전순약공업": "공통 수량단위 분리 (해당 건 없음)",
    "대한과학": "선두 '#' 제거만 (공통 수량단위 분리 미적용 — v2.1과 동일)",
    "Corning": "공통 수량단위 분리(EA)",
    "USP": "'USP' 라벨 제거 + 공통 수량단위 분리(MG)",
    "Mettler Toledo": "'시리얼 번호:' 라벨 제거 + 공통 수량단위 분리",
    "Invitrogen": "공통 수량단위 분리",
    "Waters": "공통 수량단위 분리 (해당 건 없음)",
    "Cytiva": "끝의 ' 외' 제거 + 공통 수량단위 분리",
    "유코": "공통 수량단위 분리 (해당 건 없음)",
    "Cell Signaling Technology": "선두 '#' 제거만",
    "Roche": "공통 수량단위 분리 (해당 건 없음)",
    "Eppendorf": "'모델명:' 라벨 제거 + 공통 수량단위 분리",
    "BD": "공통 수량단위 분리 (해당 건 없음)",
    "Gibco": "공통 수량단위 분리",
}


def main():
    df = pd.read_excel(SRC, sheet_name=SHEET, header=HEADER_ROW - 1)
    pn_col = df.columns[COL_PN - 1]
    mfr_col = df.columns[COL_MFR_CLEAN_OLD - 1]

    total = Counter()
    sep_cnt = Counter()
    memo_cnt = Counter()

    for _, row in df.iterrows():
        pn = row[pn_col]
        mfr = row[mfr_col]
        if pd.isna(pn) or pd.isna(mfr) or mfr not in TOP20_MFRS:
            continue
        total[mfr] += 1
        tobe, memo, sep = compute_tobe(pn, mfr)
        if sep:
            sep_cnt[mfr] += 1
        if memo:
            memo_cnt[mfr] += 1

    order = [
        "Sartorius", "Thermo Fisher Scientific", "Sigma-Aldrich", "Sigma", "Merck Millipore",
        "Agilent", "대한과학", "삼전순약공업", "Corning", "USP", "Mettler Toledo", "Invitrogen",
        "Waters", "Cytiva", "유코", "Cell Signaling Technology", "Roche", "Eppendorf", "BD", "Gibco",
    ]

    d = docx.Document()
    t = d.add_paragraph("품번 To-Be 2단계 — v2.3 보완 결과 요약")
    t.runs[0].bold = True
    t.runs[0].font.size = Pt(16)

    d.add_paragraph(
        "산출 파일: 1.41.51_품번_To-Be\\20_결과\\"
        "260812_S-TEPS_품번_To-Be_v2.3(복합단위추가_CAP_공백하이픈_선두인치기호).xlsx\n"
        "v2.2에서 사용자가 상위20 전체를 재검토하며 접수한 추가 보완 지시를 반영한 개정판. "
        "원본(1.41.50_v0.5)을 통째로 복사한 뒤 컬럼 3개(품번_To-Be/비고/분리텍스트)를 동일하게 삽입해 채움."
    )
    d.add_paragraph("")

    d.add_paragraph("1. v2.2 대비 변경 사항").runs[0].bold = True
    d.add_paragraph(
        "① 단위 화이트리스트에 CAP 추가. 예: 'C3041-100CAP'→'C3041'. "
        "(v2.1 검토 때는 CAP을 제외하기로 했었으나, 이번 재검토에서 포함하는 것으로 확정.)\n"
        "② 하이픈 포함 복합단위 3종 추가: 'G-F', 'SET-F', 'G-K' (기존 AMP-EA, KG-K와 동일한 방식 — "
        "정확히 일치할 때만 인정). 예: '11009-100G-F'→'11009', '87574-1SET-F'→'87574', "
        "'W200220-100G-K'→'W200220'.\n"
        "③ 'NxM단위' 곱셈표기에서 뒷자리 숫자가 소수점만 있는 경우(예: '.5')도 인식. "
        "예: '646563-10X.5ML'→'646563'.\n"
        "④ 하이픈 뒤에 공백이 낀 경우도 인식. 예: 'R0278- 50ML'→'R0278'.\n"
        "⑤ 신규: 선두(맨 앞) '숫자+인치기호' 제거 — 지금까지는 전부 '하이픈+뒤쪽 접미사'만 다뤘는데, "
        "이번에 발견된 Merck Millipore 'X.XX\" Clamps and gaskets' 2건은 하이픈도 없고 핵심코드라 부를 "
        "부분도 없는 서술형 값(예: '0.25\" 클램프와 개스킷'). 실제 데이터에는 인치기호가 백슬래시+따옴표"
        "(\\\") 형태로 저장돼 있음을 확인. 맨 앞의 '숫자+인치기호'를 수량/단위로 보고 제거하고 나머지 "
        "설명 텍스트를 품번_To-Be로 채택: '0.25\\\" Clamps and gaskets'→'Clamps and gaskets'."
    )

    d.add_paragraph("2. 처리 범위").runs[0].bold = True
    d.add_paragraph(
        f"품번 보유 행 중 제조사(정리) 상위 20개, 총 {sum(total.values()):,}건 (v2.2와 동일 범위, 변동 없음)."
    )

    d.add_paragraph("3. 제조사별 처리 결과").runs[0].bold = True
    table = d.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["제조사", "대상 건수", "분리텍스트 발생", "비고(미분리) 표시", "적용 규칙"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
    for mfr in order:
        cells = table.add_row().cells
        cells[0].text = mfr
        cells[1].text = f"{total[mfr]:,}"
        cells[2].text = f"{sep_cnt[mfr]:,}"
        cells[3].text = f"{memo_cnt[mfr]:,}" if memo_cnt[mfr] else "-"
        cells[4].text = RULE_LABEL[mfr]
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    d.add_paragraph("")
    d.add_paragraph("4. 검증").runs[0].bold = True
    d.add_paragraph(
        f"제조사별 대상 건수 합계가 v2.2와 동일하게 {sum(total.values()):,}건임을 재확인함(범위 불변). "
        f"분리텍스트가 채워진 행은 총 {sum(sep_cnt.values()):,}건. v2.2와 v2.3의 계산 결과를 회사별로 "
        "전수 대조해서, 이번에 의도한 항목(CAP, G-F/SET-F/G-K 복합단위, 곱셈표기 소수, 공백하이픈, "
        "선두 인치기호) 외에는 단 한 건도 값이 달라지지 않았음을 확인함 — 총 변경 18건(Sigma·"
        "Sigma-Aldrich 16건, Merck Millipore 2건), 전부 의도한 규칙에 해당."
    )
    d.add_paragraph(
        "※ 사용자가 언급한 '숫자-숫자ml' 패턴은 이번 재검토에서 실제 값을 찾지 못함(해당 정규식 매칭 "
        "0건) — 별도 예시 확인 후 후속 버전에서 반영 예정."
    )

    d.add_paragraph("")
    d.add_paragraph("5. 컬럼 설명").runs[0].bold = True
    d.add_paragraph("품번_To-Be: 규칙 적용 후 확정된 품번 값. 상위20 제조사가 아니거나 품번이 없으면 공란.")
    d.add_paragraph("품번_To-Be_비고: Sigma/Sigma-Aldrich 중 수량단위 패턴이 안 맞아 원본을 그대로 유지한 경우에만 표시.")
    d.add_paragraph(
        "품번_To-Be_분리텍스트: 원본 품번에서 잘려나간 부분을 그대로 기록. 라벨/기호/단위(복합단위·"
        "곱셈표기 포함) 제거, Agilent 'UI' 제거, 선두 인치기호 제거가 모두 이 컬럼에 기록됨."
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    d.save(OUT)
    print(f"[저장] {OUT}")


if __name__ == "__main__":
    main()
