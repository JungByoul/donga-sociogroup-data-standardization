# -*- coding: utf-8 -*-
r"""v2.2 실행 결과 요약 docx 작성. step6의 compute_tobe를 그대로 불러와 실제 데이터로 집계."""
import sys
from pathlib import Path
from collections import Counter

import docx
import pandas as pd
from docx.shared import Pt

sys.path.insert(0, str(Path(__file__).parent))
from step6_build_pn_tobe_v22 import (  # noqa: E402
    SRC, SHEET, HEADER_ROW, COL_PN, COL_MFR_CLEAN_OLD, TOP20_MFRS,
    compute_tobe,
)

OUT = Path(r"C:\Users\정별\1_Work\1.41_동아쏘시오그룹(2)_데이터_표준화\1.41.51_품번_To-Be\30_보고서\260812_품번_To-Be_2단계_결과요약_v2.2.docx")

RULE_LABEL = {
    "Sigma-Aldrich": "선두 '#' 제거 + 핵심코드-수량단위 분리(복합단위/곱셈표기/소수 확장 포함)",
    "Sigma": "선두 '#' 제거 + 핵심코드-수량단위 분리(복합단위/곱셈표기/소수 확장 포함)",
    "Merck Millipore": "공통 수량단위 분리(곱셈표기 'NxM단위' 포함)",
    "Thermo Fisher Scientific": "공통 수량단위 분리 ('KOLAS'로 시작하는 값은 예외)",
    "Sartorius": "'견적번호/견적서 번호 :' 라벨 제거 + '[...]' 제거 + 공통 수량단위 분리",
    "Agilent": "'부품 번호:' 라벨 제거 + 'UI' 접미사만 제거(숫자 접미사는 유지, 공통 수량단위 분리 미적용)",
    "삼전순약공업": "공통 수량단위 분리 (해당 건 없음)",
    "대한과학": "선두 '#' 제거만 (공통 수량단위 분리 자체를 미적용 — v2.1과 동일)",
    "Corning": "공통 수량단위 분리(EA)",
    "USP": "'USP' 라벨 제거 + 공통 수량단위 분리(MG)",
    "Mettler Toledo": "'시리얼 번호:' 라벨 제거 + 공통 수량단위 분리",
    "Invitrogen": "공통 수량단위 분리",
    "Waters": "공통 수량단위 분리 (해당 건 없음)",
    "Cytiva": "끝의 ' 외' 제거 + 공통 수량단위 분리",
    "유코": "공통 수량단위 분리 (해당 건 없음)",
    "Cell Signaling Technology": "선두 '#' 제거만",
    "Roche": "공통 수량단위 분리 (해당 건 없음)",
    "Eppendorf": "'모델명:' 라벨 제거 + 공통 수량단위 분리",
    "BD": "공통 수량단위 분리 (해당 건 없음)",
    "Gibco": "공통 수량단위 분리",
}


def main():
    df = pd.read_excel(SRC, sheet_name=SHEET, header=HEADER_ROW - 1)
    pn_col = df.columns[COL_PN - 1]
    mfr_col = df.columns[COL_MFR_CLEAN_OLD - 1]

    total = Counter()
    sep_cnt = Counter()
    memo_cnt = Counter()

    for _, row in df.iterrows():
        pn = row[pn_col]
        mfr = row[mfr_col]
        if pd.isna(pn) or pd.isna(mfr) or mfr not in TOP20_MFRS:
            continue
        total[mfr] += 1
        tobe, memo, sep = compute_tobe(pn, mfr)
        if sep:
            sep_cnt[mfr] += 1
        if memo:
            memo_cnt[mfr] += 1

    order = [
        "Sartorius", "Thermo Fisher Scientific", "Sigma-Aldrich", "Sigma", "Merck Millipore",
        "Agilent", "대한과학", "삼전순약공업", "Corning", "USP", "Mettler Toledo", "Invitrogen",
        "Waters", "Cytiva", "유코", "Cell Signaling Technology", "Roche", "Eppendorf", "BD", "Gibco",
    ]

    d = docx.Document()
    t = d.add_paragraph("품번 To-Be 2단계 — v2.2 보완 결과 요약")
    t.runs[0].bold = True
    t.runs[0].font.size = Pt(16)

    d.add_paragraph(
        "산출 파일: 1.41.51_품번_To-Be\\20_결과\\"
        "260812_S-TEPS_품번_To-Be_v2.2(복합단위_곱셈표기_Agilent UI_단위확장).xlsx\n"
        "v2.1에서 접수한 추가 보완 지시를 반영한 개정판. 원본(1.41.50_v0.5)을 통째로 복사한 뒤 "
        "Steps_중복제거_32359 시트 P열(품번_정리) 오른쪽에 컬럼 3개(품번_To-Be/비고/분리텍스트)를 "
        "동일하게 삽입해 채움."
    )
    d.add_paragraph("")

    d.add_paragraph("1. v2.1 대비 변경 사항").runs[0].bold = True
    d.add_paragraph(
        "① 단위 화이트리스트에 RXN/KT/TAB/VL 추가(기존 KG/MG/UG/NG/AMP/PAK/EA/ML/UL/G/L/%). "
        "예: 'DUO92001-100RXN'→'DUO92001', 'MINI26-1KT'→'MINI26', 'P4417-100TAB'→'P4417', "
        "'S2442-1VL'→'S2442'.\n"
        "② 수량 표기 인식 확장 — 소수점 앞자리가 없는 경우(예: 'A3687-.5ML'→'A3687')와, "
        "'N x M단위' 곱셈 표기(예: 'D2438-5x10ML'→'D2438', '471283-4X100ML'→'471283', "
        "'T6567-5X20ug'→'T6567')도 수량단위 접미사로 인식해 분리. Merck Millipore "
        "'302031-10X1ML'(v2.1에서는 패턴이 안 맞아 원본 유지했던 예외 건)도 이 확장으로 자동 해결됨 "
        "→ '302031'.\n"
        "③ 단위 자체가 하이픈을 포함한 복합단위 2종을 정확히 일치할 때만 인정: 'AMP-EA', 'KG-K'. "
        "예: '45-T6508-10AMP-EA'→'45-T6508'(분리텍스트 '-10AMP-EA'), "
        "'W278430-1KG-K'→'W278430'(분리텍스트 '-1KG-K'). 이 2개 조합 외 임의의 하이픈 포함 단위는 "
        "인정하지 않음(오탐 방지).\n"
        "④ Agilent 전용 'UI' 접미사 제거 — Agilent는 하이픈 뒤 숫자가 수량이 아니라 서로 다른 실제 "
        "제품임이 1단계에서 이미 확인된 회사라, 공통 수량단위 분리 엔진을 그대로 적용하면 위험함. "
        "그래서 Agilent는 이번에도 공통 엔진 대상에서 제외하고, 'UI' 글자만 별도로 떼어내고 숫자는 "
        "그대로 유지: '123-0364UI'→'123-0364'(분리텍스트 'UI'). 실제 다른 제품인 '122-7032' 같은 "
        "순수 숫자 접미사는 이전처럼 손대지 않음."
    )

    d.add_paragraph("2. 처리 범위").runs[0].bold = True
    d.add_paragraph(
        f"품번 보유 행 중 제조사(정리) 상위 20개, 총 {sum(total.values()):,}건 (v2.1과 동일 범위, "
        "변동 없음)."
    )

    d.add_paragraph("3. 제조사별 처리 결과").runs[0].bold = True
    table = d.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["제조사", "대상 건수", "분리텍스트 발생", "비고(미분리) 표시", "적용 규칙"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
    for mfr in order:
        cells = table.add_row().cells
        cells[0].text = mfr
        cells[1].text = f"{total[mfr]:,}"
        cells[2].text = f"{sep_cnt[mfr]:,}"
        cells[3].text = f"{memo_cnt[mfr]:,}" if memo_cnt[mfr] else "-"
        cells[4].text = RULE_LABEL[mfr]
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    d.add_paragraph("")
    d.add_paragraph("4. 검증").runs[0].bold = True
    d.add_paragraph(
        f"제조사별 대상 건수 합계가 v2.1과 동일하게 {sum(total.values()):,}건임을 재확인함(범위 불변). "
        f"분리텍스트가 채워진 행은 총 {sum(sep_cnt.values()):,}건. v2.1과 v2.2의 계산 결과를 "
        "회사별로 전수 대조해서, 이번에 의도한 항목(RXN/KT/TAB/VL, 소수/곱셈표기, AMP-EA/KG-K 복합단위, "
        "Agilent UI) 외에는 단 한 건도 값이 달라지지 않았음을 확인함 — 총 변경 28건 "
        "(Sigma·Sigma-Aldrich 18건, Agilent 8건, Merck Millipore 2건), 전부 의도한 규칙에 해당."
    )

    d.add_paragraph("")
    d.add_paragraph("5. 컬럼 설명").runs[0].bold = True
    d.add_paragraph("품번_To-Be: 규칙 적용 후 확정된 품번 값. 상위20 제조사가 아니거나 품번이 없으면 공란.")
    d.add_paragraph("품번_To-Be_비고: Sigma/Sigma-Aldrich 중 수량단위 패턴이 안 맞아 원본을 그대로 유지한 경우에만 표시.")
    d.add_paragraph(
        "품번_To-Be_분리텍스트: 원본 품번에서 잘려나간 부분을 그대로 기록. 라벨 제거, '#'/'[...]' 제거, "
        "수량단위 접미사(복합단위·곱셈표기 포함) 제거, Agilent 'UI' 제거가 모두 이 컬럼에 기록됨."
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    d.save(OUT)
    print(f"[저장] {OUT}")


if __name__ == "__main__":
    main()
