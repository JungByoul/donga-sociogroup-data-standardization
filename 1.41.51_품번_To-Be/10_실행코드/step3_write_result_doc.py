# -*- coding: utf-8 -*-
r"""2단계 실행 결과 요약 docx 작성."""
from pathlib import Path

import docx
from docx.shared import Pt

OUT = Path(r"C:\Users\정별\1_Work\1.41_동아쏘시오그룹(2)_데이터_표준화\1.41.51_품번_To-Be\30_보고서\260811_품번_To-Be_2단계_결과요약_v2.0.docx")

ROWS = [
    ("Sigma-Aldrich", "615", "218", "핵심코드-수량단위 분리", "397건은 패턴 불일치로 원본 트림값 유지"),
    ("Sigma", "500", "111", "핵심코드-수량단위 분리", "389건은 패턴 불일치로 원본 트림값 유지"),
    ("Merck Millipore", "436", "436", "트림만", "-"),
    ("Thermo Fisher Scientific", "631", "631", "트림만", "-"),
    ("Sartorius", "607", "607", "트림 + '[...]' 제거", "-"),
    ("Agilent", "274", "274", "트림만(접미사 분리 안 함)", "-"),
    ("삼전순약공업", "247", "247", "트림만", "-"),
    ("대한과학", "194", "194", "트림 + 선두 '#' 제거", "-"),
    ("Corning", "180", "180", "트림만(접미사 분리 안 함)", "-"),
    ("USP", "177", "177", "트림만", "-"),
    ("Mettler Toledo", "161", "161", "트림만", "-"),
    ("Invitrogen", "154", "154", "트림만", "-"),
    ("Waters", "140", "140", "트림만", "-"),
    ("Cytiva", "138", "138", "트림만", "-"),
    ("유코", "120", "120", "트림만", "-"),
    ("Cell Signaling Technology", "119", "119", "트림 + 선두 '#' 제거", "-"),
    ("Roche", "113", "113", "트림만", "-"),
    ("Eppendorf", "101", "101", "트림만", "-"),
    ("BD", "98", "98", "트림만", "-"),
    ("Gibco", "97", "97", "트림만", "-"),
]


def main():
    d = docx.Document()
    t = d.add_paragraph("품번 To-Be 2단계 — 컬럼 생성 결과 요약")
    t.runs[0].bold = True
    t.runs[0].font.size = Pt(16)

    d.add_paragraph(
        "산출 파일: 1.41.51_품번_To-Be\\20_결과\\260811_S-TEPS_품번_To-Be_v2.0(원본유지).xlsx\n"
        "원본(1.41.50_v0.5) 파일을 통째로 복사해서(시트 7개·서식·다른 시트 전부 그대로 유지) "
        "그 복사본의 Steps_중복제거_32359 시트에 실제로 P열(품번_정리) 오른쪽 컬럼 2개를 삽입해서 채움. "
        "(※ 최초 v1.0은 시트 하나만 값으로 새로 만드는 방식으로 시도했으나, 컬럼 밀림 처리가 잘못돼 "
        "옆 컬럼 데이터와 뒤섞이는 오류가 있어 폐기하고 이 방식으로 다시 만듦)\n"
        "컬럼 삽입 시, 이 시트 안의 유일한 수식(M2: =COUNTIF(AI6:AI32354,\"*!*\"))의 참조도 "
        "AK6:AK32354로 함께 보정함. 다른 시트가 이 시트를 참조하는 수식은 없음을 사전 확인함."
    )
    d.add_paragraph("")

    d.add_paragraph("1. 처리 범위").runs[0].bold = True
    d.add_paragraph(
        "품번 보유 행(14,013건) 중 제조사(정리) 상위 20개, 총 5,102건(36.4%)에 대해서만 처리. "
        "나머지 행은 이번 단계 범위 밖으로 공란 유지."
    )

    d.add_paragraph("2. 적용 규칙 상세").runs[0].bold = True
    d.add_paragraph(
        "① 핵심코드-수량단위 분리 (Sigma, Sigma-Aldrich 전용) — 원본 품번에서 마지막 하이픈(-) "
        "뒤에 오는 부분이 '숫자(소수점 가능)+단위' 형태이고, 그 단위가 화이트리스트(G/KG/MG/UG/NG/"
        "ML/L/UL/%)와 정확히 일치할 때만 그 부분을 잘라내고 앞부분만 품번_To-Be로 채택. "
        "예: '09735-250G' → '09735'(250G 제거). 하이픈+숫자 패턴이 있어도 단위가 화이트리스트에 "
        "없으면(예: 'DUO92001-100RXN'의 'RXN') 분리하지 않고 트림한 원본을 그대로 유지하며, "
        "품번_To-Be_비고에 '핵심코드 분리 안 됨(수량단위 패턴 불일치) - 원본 트림값 유지'라고 표시.\n"
        "② 선두 '#' 제거 (대한과학, Cell Signaling Technology 전용) — 값이 '#'로 시작하는 경우에만 "
        "맨 앞 '#' 한 글자를 제거. 예: '#110081' → '110081'. 하이픈 뒤 접미사는 건드리지 않음 "
        "(예: '#2319-0050' → '2319-0050', 뒤의 '-0050'은 그대로 유지).\n"
        "③ '[...]' 괄호 제거 (Sartorius 전용) — 값 전체가 대괄호로 감싸진 경우에만(시작이 '[', 끝이 "
        "']') 양쪽 대괄호만 제거. 예: '[16421-E]' → '16421-E'. 부분적으로만 괄호가 있는 경우는 대상 "
        "아님.\n"
        "④ 나머지 16개 회사(Merck Millipore, Thermo Fisher Scientific, Agilent, 삼전순약공업, "
        "Corning, USP, Mettler Toledo, Invitrogen, Waters, Cytiva, 유코, Roche, Eppendorf, BD, "
        "Gibco) — 트림(앞뒤 공백 제거)만 적용, 구조는 바꾸지 않음. 특히 Agilent·Corning은 1단계 "
        "패턴분석 때 하이픈 뒤 숫자가 수량이 아니라 서로 다른 실제 제품임을 원문 대조로 확인했기 "
        "때문에(예: Agilent '122-0733~122-7032', Corning 'COP-430281~430291'는 각각 다른 제품) "
        "접미사 제거 로직 자체를 의도적으로 적용하지 않음."
    )

    d.add_paragraph("3. 제조사별 처리 결과").runs[0].bold = True
    table = d.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["제조사", "대상 건수", "To-Be 채움", "적용 규칙", "비고"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
    for mfr, total, filled, rule, note in ROWS:
        cells = table.add_row().cells
        cells[0].text = mfr
        cells[1].text = total
        cells[2].text = filled
        cells[3].text = rule
        cells[4].text = note
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    d.add_paragraph("")
    d.add_paragraph("4. 검증").runs[0].bold = True
    d.add_paragraph(
        "제조사별 To-Be 채움 건수를 1단계 분석 수치와 재대조해서 전부 일치함을 확인함(합계 5,102건). "
        "Sigma·Sigma-Aldrich 핵심코드 분리 성공 329건 / 실패(원본 유지) 786건도 원본 품번·품목명과 "
        "나란히 놓고 육안 샘플 검토를 마침(예: 09735-250G → 09735, Ammonium formate)."
    )
    d.add_paragraph("")
    d.add_paragraph("5. 컬럼 설명").runs[0].bold = True
    d.add_paragraph("품번_To-Be: 위 규칙으로 정규화된 확정 품번 값. 상위20 제조사가 아니거나 품번이 없으면 공란.")
    d.add_paragraph("품번_To-Be_비고: 규칙이 있었지만(Sigma류) 패턴이 안 맞아 원본을 그대로 유지한 경우에만 표시.")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    d.save(OUT)
    print(f"[저장] {OUT}")


if __name__ == "__main__":
    main()
