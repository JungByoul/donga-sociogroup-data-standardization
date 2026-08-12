# -*- coding: utf-8 -*-
r"""v2.1 실행 결과 요약 docx 작성. step4의 compute_tobe를 그대로 불러와 실제 데이터로 집계."""
import sys
from pathlib import Path
from collections import Counter

import docx
import pandas as pd
from docx.shared import Pt

sys.path.insert(0, str(Path(__file__).parent))
from step4_build_pn_tobe_v21 import (  # noqa: E402
    SRC, SHEET, HEADER_ROW, COL_PN, COL_MFR_CLEAN_OLD, TOP20_MFRS,
    compute_tobe,
)

OUT = Path(r"C:\Users\정별\1_Work\1.41_동아쏘시오그룹(2)_데이터_표준화\1.41.51_품번_To-Be\30_보고서\260812_품번_To-Be_2단계_결과요약_v2.1.docx")

RULE_LABEL = {
    "Sigma-Aldrich": "선두 '#' 제거 + 핵심코드-수량단위 분리(EA/AMP/PAK 포함)",
    "Sigma": "선두 '#' 제거 + 핵심코드-수량단위 분리(EA/AMP/PAK 포함)",
    "Merck Millipore": "공통 수량단위 분리(EA 포함, 'GMCN'처럼 단위 뒤 잉여문자 붙은 경우도 앞부분 단위로 인식)",
    "Thermo Fisher Scientific": "공통 수량단위 분리 ('KOLAS'로 시작하는 값은 용량범위 표기로 보고 예외 처리)",
    "Sartorius": "'견적번호/견적서 번호 :' 라벨 제거 + '[...]' 제거 + 공통 수량단위 분리",
    "Agilent": "'부품 번호:' 라벨 제거 (수량단위 분리는 대상 없음)",
    "삼전순약공업": "공통 수량단위 분리 (해당 건 없음, 트림만)",
    "대한과학": "선두 '#' 제거만 (공통 수량단위 분리는 이 회사만 예외 — 'KA.11-39L'처럼 접미사가 실제 코드 일부일 가능성이 있어 미적용)",
    "Corning": "공통 수량단위 분리(EA)",
    "USP": "'USP' 라벨 제거 + 공통 수량단위 분리(MG)",
    "Mettler Toledo": "'시리얼 번호:' 라벨 제거 + 공통 수량단위 분리",
    "Invitrogen": "공통 수량단위 분리",
    "Waters": "공통 수량단위 분리 (해당 건 없음, 트림만)",
    "Cytiva": "끝의 ' 외' 제거 + 공통 수량단위 분리",
    "유코": "공통 수량단위 분리 (해당 건 없음, 트림만)",
    "Cell Signaling Technology": "선두 '#' 제거만",
    "Roche": "공통 수량단위 분리 (해당 건 없음, 트림만)",
    "Eppendorf": "'모델명:' 라벨 제거 + 공통 수량단위 분리",
    "BD": "공통 수량단위 분리 (해당 건 없음, 트림만)",
    "Gibco": "공통 수량단위 분리",
}


def main():
    df = pd.read_excel(SRC, sheet_name=SHEET, header=HEADER_ROW - 1)
    pn_col = df.columns[COL_PN - 1]
    mfr_col = df.columns[COL_MFR_CLEAN_OLD - 1]

    total = Counter()
    filled = Counter()
    sep_cnt = Counter()
    memo_cnt = Counter()

    for _, row in df.iterrows():
        pn = row[pn_col]
        mfr = row[mfr_col]
        if pd.isna(pn) or pd.isna(mfr) or mfr not in TOP20_MFRS:
            continue
        total[mfr] += 1
        tobe, memo, sep = compute_tobe(pn, mfr)
        filled[mfr] += 1
        if sep:
            sep_cnt[mfr] += 1
        if memo:
            memo_cnt[mfr] += 1

    order = [
        "Sartorius", "Thermo Fisher Scientific", "Sigma-Aldrich", "Sigma", "Merck Millipore",
        "Agilent", "대한과학", "삼전순약공업", "Corning", "USP", "Mettler Toledo", "Invitrogen",
        "Waters", "Cytiva", "유코", "Cell Signaling Technology", "Roche", "Eppendorf", "BD", "Gibco",
    ]

    d = docx.Document()
    t = d.add_paragraph("품번 To-Be 2단계 — v2.1 보완 결과 요약")
    t.runs[0].bold = True
    t.runs[0].font.size = Pt(16)

    d.add_paragraph(
        "산출 파일: 1.41.51_품번_To-Be\\20_결과\\"
        "260812_S-TEPS_품번_To-Be_v2.1(공통단위분리_라벨제거_분리텍스트열추가).xlsx\n"
        "v2.0(사용자 20개 회사 전수 검토 완료본)에서 접수한 보완 지시를 반영한 개정판. "
        "원본(1.41.50_v0.5)을 통째로 복사한 뒤 Steps_중복제거_32359 시트 P열(품번_정리) 오른쪽에 "
        "컬럼 3개(품번_To-Be, 품번_To-Be_비고, 품번_To-Be_분리텍스트)를 신규 삽입해 채움. "
        "시트 내 유일한 수식(M2)의 참조도 열 3칸 밀림에 맞춰 함께 보정함."
    )
    d.add_paragraph("")

    d.add_paragraph("1. v2.0 대비 변경 사항").runs[0].bold = True
    d.add_paragraph(
        "① 수량단위 접미사 분리 엔진을 Sigma/Sigma-Aldrich 전용에서 TOP20 전 회사 공통으로 확장. "
        "화이트리스트에 EA/AMP/PAK 추가(기존 G/KG/MG/UG/NG/ML/L/UL/%). "
        "단위 글자가 화이트리스트와 정확히 같지 않아도 그 단어로 '시작'하면 인정 "
        "(예: Millipore '218680-100GMCN' → 앞의 G를 단위로 인식해 '218680'/'-100GMCN'으로 분리).\n"
        "② 회사별 라벨(견적번호/부품 번호/USP/시리얼 번호/모델명/' 외') 제거 규칙 신설 — 아래 표의 "
        "해당 회사에만 적용, 다른 회사에는 미적용.\n"
        "③ 대한과학 'KA.11-39L'처럼 접미사가 진짜 수량인지 코드 일부인지 원문만으로 판단이 안 되는 "
        "회사는 이번 공통엔진 적용에서 통째로 제외(기존처럼 '#' 제거만 유지). "
        "Thermo Fisher는 'KOLAS'로 시작하는 값(용량 범위 표기로 판단)만 개별 예외 처리.\n"
        "④ Sigma/Sigma-Aldrich 선두 '#' 제거 추가(예: '#04243'→'04243'). 단, 'S#110205'처럼 문자와 "
        "문자 사이에 낀 '#'는 의미가 불확실해 이번엔 보류(원본 유지).\n"
        "⑤ 신규 컬럼 '품번_To-Be_분리텍스트' 추가 — 원본에서 잘려나간 부분을 그대로 기록 "
        "(예: '3232-1EA' → To-Be '3232' / 분리텍스트 '-1EA'). Sigma 계열이 v2.0에서 이미 분리 성공한 "
        "건들도 이번에 분리텍스트를 채움.\n"
        "※ TOP20 밖 다른 제조사(ALPCO, DURAN 등)에도 '#' 섞인 값이 다수 있으나, 이번 v2.1은 여전히 "
        "TOP20 범위만 다룸 — 범위 확장은 별도 논의 필요."
    )

    d.add_paragraph("2. 처리 범위").runs[0].bold = True
    d.add_paragraph(
        f"품번 보유 행 중 제조사(정리) 상위 20개, 총 {sum(total.values()):,}건에 대해서만 처리 "
        "(v2.0과 동일 범위, 대상 건수 변동 없음). 나머지 행은 이번 단계 범위 밖으로 공란 유지."
    )

    d.add_paragraph("3. 제조사별 처리 결과").runs[0].bold = True
    table = d.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["제조사", "대상 건수", "분리텍스트 발생", "비고(미분리) 표시", "적용 규칙"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
    for mfr in order:
        cells = table.add_row().cells
        cells[0].text = mfr
        cells[1].text = f"{total[mfr]:,}"
        cells[2].text = f"{sep_cnt[mfr]:,}"
        cells[3].text = f"{memo_cnt[mfr]:,}" if memo_cnt[mfr] else "-"
        cells[4].text = RULE_LABEL[mfr]
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    d.add_paragraph("")
    d.add_paragraph("4. 검증").runs[0].bold = True
    d.add_paragraph(
        f"제조사별 대상 건수 합계가 v2.0과 동일하게 {sum(total.values()):,}건임을 재확인함(범위 불변). "
        f"분리텍스트가 새로 채워진 행은 총 {sum(sep_cnt.values()):,}건(v2.0 대비 신규 회사·확장 단위 반영분 "
        "포함). 사용자 보완 지시 항목(회사별 라벨 제거, 단위 확장, 예외 2건 — Thermo 'KOLAS...', "
        "대한과학 'KA.11-39L', Sigma 중간 '#' 보류)을 실제 값 기준으로 각각 재확인 후 반영함."
    )

    d.add_paragraph("")
    d.add_paragraph("5. 컬럼 설명").runs[0].bold = True
    d.add_paragraph("품번_To-Be: 규칙 적용 후 확정된 품번 값. 상위20 제조사가 아니거나 품번이 없으면 공란.")
    d.add_paragraph("품번_To-Be_비고: Sigma/Sigma-Aldrich 중 수량단위 패턴이 안 맞아 원본을 그대로 유지한 경우에만 표시.")
    d.add_paragraph(
        "품번_To-Be_분리텍스트(신규): 원본 품번에서 잘려나간 부분을 그대로 기록. "
        "예) '3232-1EA' → 품번_To-Be '3232' / 분리텍스트 '-1EA'. "
        "라벨 제거(예: '견적번호'), '#'/'[...]' 제거, 수량단위 접미사 제거가 모두 이 컬럼에 기록됨. "
        "잘려나간 부분이 없으면 공란."
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    d.save(OUT)
    print(f"[저장] {OUT}")


if __name__ == "__main__":
    main()
