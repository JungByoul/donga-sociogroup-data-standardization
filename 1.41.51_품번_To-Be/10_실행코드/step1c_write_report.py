# -*- coding: utf-8 -*-
r"""
1단계 분석(제조사별 품번 패턴) 결과를 docx 보고서로 정리.
아직 규칙 확정/실행 전 단계이며, 이 문서는 분석 결과 공유 및 규칙 확정용.
"""
from pathlib import Path

import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(r"C:\Users\정별\1_Work\1.41_동아쏘시오그룹(2)_데이터_표준화\1.41.51_품번_To-Be\30_보고서\260811_품번_To-Be_1단계_제조사별_패턴분석_v1.0.docx")

ROWS = [
    # 제조사, 건수, 지배패턴(커버율), 제안 규칙, 근거/비고
    ("Sigma-Aldrich", "615", "999999 / A9999 (43.7%)", "핵심코드-수량단위 분리 (단위=G/KG/MG/ML/L 등 도량형일 때만)", "core='C7519'에서 250ML/100ML 등 실제 용량 차이 확인(8개 코드)"),
    ("Sigma", "500", "A9999 / 999999 (62.0%)", "위와 동일 규칙", "위와 동일"),
    ("Merck Millipore", "436", "9.99999.9999 (61.5%)", "접미사 제거 없음, 점(.) 구분 포맷만 통일", "기존 패턴분석(Merck_점구분_교차검증)에서 Merck 고유 카탈로그 체계로 이미 검증됨"),
    ("Thermo Fisher Scientific", "631", "분산 (상위3 32.6%)", "트림/공백 정리만, 자동 확정 보류", "지배 패턴이 약해 통일 규칙 위험"),
    ("Sartorius", "607", "AAA999999 (33.9%)", "트림 + '[...]' 감싸짐 제거(2건)", "하이픈이 코드 구조 자체의 일부(예: 001-2A23) → 접미사 제거 금지"),
    ("Agilent", "274", "9999-9999 (63.9%)", "트림만, 접미사 제거 금지", "'122-0733~122-7032'는 수량이 아니라 서로 다른 실제 제품임을 원문 대조로 확인(오탐)"),
    ("삼전순약공업", "247", "A9999 (99.6%)", "트림만", "지배율 매우 높아 이미 정제된 상태"),
    ("대한과학", "194", "분산 (24.7%)", "트림 + 선두 '#' 기호 제거", "지배 패턴 약함 → 자동확정 최소화. '#' 제거 시 중복 1건 확인"),
    ("Corning", "180", "999999 (65.0%)", "트림만, 접미사 제거 금지", "'COP-430281~430291'도 서로 다른 실제 제품임을 원문 대조로 확인(오탐)"),
    ("USP", "177", "9999999 (98.3%)", "트림만", "이미 정제된 상태"),
    ("Mettler Toledo", "161", "99999999 (85.7%)", "트림만", "-"),
    ("Invitrogen", "154", "분산 (42.2%)", "트림만, 자동확정 최소화", "지배 패턴 약함"),
    ("Waters", "140", "999999999 (96.4%)", "트림만", "이미 정제된 상태"),
    ("Cytiva", "138", "99999999 (79.0%)", "트림만", "-"),
    ("유코", "120", "99999999 (74.2%)", "트림만", "-"),
    ("Cell Signaling Technology", "119", "9999A (88.2%)", "트림 + 선두 '#' 기호 제거", "-"),
    ("Roche", "113", "11자리 숫자 (97.3%)", "트림만", "이미 정제된 상태"),
    ("Eppendorf", "101", "9999999999 (91.1%)", "트림만, 공백 포함 형태 유지", "공백 구분이 실제 표기 관례로 보임"),
    ("BD", "98", "999999 (99.0%)", "트림만", "이미 정제된 상태"),
    ("Gibco", "97", "분산 (82.5%)", "트림만, 접미사 제거 보류", "접미사매칭 40.2%지만 근거 코드 5개뿐 → 확정 보류"),
]


def set_col_widths(table, widths_cm):
    for row in table.rows:
        for cell, w in zip(row.cells, widths_cm):
            cell.width = Cm(w)


def main():
    d = docx.Document()

    title = d.add_paragraph("품번 To-Be 1단계 — 제조사별 품번 패턴 분석 보고")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(16)

    sub = d.add_paragraph(
        "1.41.51_품번_To-Be · 기준파일: 1.41.50_...uniq_v.0.5(0811_16시).xlsx · "
        "이 문서는 규칙 확정 전 1단계 분석 결과 공유용이며, 아직 어떤 데이터도 변경하지 않았음"
    )
    d.add_paragraph("")

    d.add_paragraph("0. 분석 개요").runs[0].bold = True
    d.add_paragraph(
        "품번이 있는 행 14,013건 중, 제조사(정리) 기준 상위 20개(5,102건, 36.4%)를 대상으로 "
        "각 제조사의 품번 값을 '형태 지문'(문자→A, 숫자→9, 그 외 기호는 유지)으로 변환해 지배 패턴을 확인했다. "
        "추가로 값이 '핵심코드-수량단위' 구조(예: 09735-250G)인지 정규식으로 걸러, 같은 핵심코드가 서로 다른 "
        "수량으로 반복 등장하는지 확인해서 접미사가 진짜 수량 구분자인지 검증했다."
    )
    d.add_paragraph("")

    d.add_paragraph("1. 핵심 발견 — '수량 접미사' 가설은 대부분 오탐").runs[0].bold = True
    d.add_paragraph(
        "1차 정규식 결과만 보면 Agilent(84.3%), Corning(35.0%) 등에서도 '핵심코드-수량' 패턴이 높게 나왔으나, "
        "원문을 직접 대조한 결과 오탐으로 확인됨:"
    )
    p = d.add_paragraph()
    p.add_run("- Agilent 예시: ").bold = True
    p.add_run("122-0733, 122-1032, 122-1334, 122-1732, 122-2232, 122-3232, 122-7032 → "
               "'122'는 제품군 접두사이고 뒤 숫자는 전부 서로 다른 실제 카탈로그 번호(수량 아님)")
    p2 = d.add_paragraph()
    p2.add_run("- Corning 예시: ").bold = True
    p2.add_run("COP-430281, COP-430282, COP-430291 ... → 전부 별개 제품(수량 아님)")
    d.add_paragraph(
        "반면 Sigma / Sigma-Aldrich는 접미사가 G·KG·MG·ML·L 등 실제 도량형 단위 문자로 끝나는 경우가 많고, "
        "같은 핵심코드가 서로 다른 용량으로 반복 등장하는 사례를 원문에서 확인함(예: core='C7519' → 250ML/100ML). "
        "따라서 접미사 제거(핵심코드 To-Be 도출)는 Sigma·Sigma-Aldrich에만 적용 근거가 있음."
    )
    d.add_paragraph("")

    d.add_paragraph("2. 제조사별 제안 규칙 (확정 전 — 검토 요청)").runs[0].bold = True
    table = d.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["제조사", "품번보유 건수", "지배 패턴(커버율)", "제안 규칙", "근거/비고"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True

    for mfr, cnt, pat, rule, note in ROWS:
        cells = table.add_row().cells
        cells[0].text = mfr
        cells[1].text = cnt
        cells[2].text = pat
        cells[3].text = rule
        cells[4].text = note

    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    d.add_paragraph("")
    d.add_paragraph("3. 요약").runs[0].bold = True
    d.add_paragraph(
        "Sigma·Sigma-Aldrich(1,115건, 전체 품번보유행의 약 8.0%, 상위20 대상의 약 21.9%)만 '핵심코드-수량 분리'로 "
        "실질적인 정규화(같은 제품의 여러 용량 주문을 하나의 코드로 수렴) 효과가 있다. "
        "나머지 18개 제조사(3,987건)는 지배 패턴 커버율이 이미 높거나(65%~99%) 낮아도(24%~43%) 뚜렷한 접미사 구조가 "
        "없어, 트림/공백정리/장식기호 제거 수준의 안전한 정리만 적용하는 것을 제안한다."
    )
    d.add_paragraph("")
    d.add_paragraph("4. 다음 단계").runs[0].bold = True
    d.add_paragraph(
        "본 문서의 제조사별 제안 규칙을 확정받은 뒤, 1.41.50 원본은 건드리지 않고 새 버전 파일을 만들어 "
        "P열(품번 정리) 오른쪽에 '품번_To-Be' 컬럼을 추가하고 규칙을 적용한다. 규칙 매칭이 안 되거나 지배 패턴에서 "
        "크게 벗어나는 값은 자동 확정하지 않고 공란 + 검토메모로 남겨 사람이 볼 수 있게 한다."
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    d.save(OUT)
    print(f"[저장] {OUT}")


if __name__ == "__main__":
    main()
