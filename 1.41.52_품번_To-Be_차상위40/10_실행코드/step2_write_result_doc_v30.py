# -*- coding: utf-8 -*-
r"""v3.0 실행 결과 요약 docx 작성. step1의 compute_tobe를 그대로 불러와 실제 데이터로 집계."""
import sys
from pathlib import Path
from collections import Counter

import docx
import pandas as pd
from docx.shared import Pt

sys.path.insert(0, str(Path(__file__).parent))
from step1_build_pn_tobe_v30 import (  # noqa: E402
    SRC, SHEET, HEADER_ROW, COL_PN, COL_MFR_CLEAN_OLD, NEW40_MFRS,
    compute_tobe,
)

OUT = Path(r"C:\Users\정별\1_Work\1.41_동아쏘시오그룹(2)_데이터_표준화\1.41.52_품번_To-Be_차상위40\30_보고서\260812_품번_To-Be_차상위40_결과요약_v3.0.docx")

ORDER = [
    "Charles River", "Km", "R&D Systems", "Combi-Blocks", "Abcam", "PerkinElmer", "대정화금",
    "BioLegend", "SPL", "Saint-Gobain", "Supelco", "싸토리우스코리아", "영인에스티", "TCI", "비티알",
    "DURAN", "Promega", "DAE", "EP", "Cobetter", "MedChemExpress", "ProteinSimple", "GEMÜ",
    "Lonza", "Metrohm", "Avantor", "Beckman Coulter", "Agilent Technologies", "(주)BB", "TRC",
    "Repligen", "Millipore", "Isolab", "Siemens", "Daejung", "Gilson", "DUKSAN", "Rhawn",
    "시너지이노베이션", "Advanced Instruments",
]


def main():
    df = pd.read_excel(SRC, sheet_name=SHEET, header=HEADER_ROW - 1)
    pn_col = df.columns[COL_PN - 1]
    mfr_col = df.columns[COL_MFR_CLEAN_OLD - 1]

    total = Counter()
    sep_cnt = Counter()

    for _, row in df.iterrows():
        pn = row[pn_col]
        mfr = row[mfr_col]
        if pd.isna(pn) or pd.isna(mfr) or mfr not in NEW40_MFRS:
            continue
        total[mfr] += 1
        tobe, memo, sep = compute_tobe(pn, mfr)
        if sep:
            sep_cnt[mfr] += 1

    d = docx.Document()
    t = d.add_paragraph("품번 To-Be — 차상위40(21~60위) v3.0 결과 요약")
    t.runs[0].bold = True
    t.runs[0].font.size = Pt(16)

    d.add_paragraph(
        "산출 파일: 1.41.52_품번_To-Be_차상위40\\20_결과\\"
        "260812_S-TEPS_품번_To-Be_v3.0(차상위40_v2.3규칙적용).xlsx\n"
        "1.41.51에서 상위20 회사를 대상으로 완성한 v2.3 로직(compute_tobe)을 규칙 변경 없이 그대로 "
        "재사용해서, 제조사(정리) 기준 21~60위(차상위40) 회사에 적용함. 원본(1.41.50_v0.5)을 통째로 "
        "복사한 뒤 동일하게 품번_To-Be/비고/분리텍스트 3개 컬럼을 삽입해 채움."
    )
    d.add_paragraph("")

    d.add_paragraph("1. 처리 범위 및 대상 선정").runs[0].bold = True
    d.add_paragraph(
        f"제조사(정리) 컬럼 기준 21~60위, 총 {sum(total.values()):,}건. "
        "1~20위(상위20, 1.41.51에서 완료)는 이번 범위에서 제외."
    )
    d.add_paragraph(
        "※ 주의사항: 21~60위 중 'Agilent Technologies'(38건), 'Millipore'(36건), 'Daejung'(36건)은 "
        "이미 처리한 'Agilent', 'Merck Millipore', 'Daejung의 한글표기 27위 대정화금'과 실제 품번이 "
        "일부 겹침을 확인함(제조사명 정리가 안 돼 같은 회사가 다른 이름으로 분리된 것으로 추정: "
        "Agilent Technologies↔Agilent 8건 겹침, Millipore↔Merck Millipore 1건 겹침, Daejung↔대정화금 "
        "2건 겹침). 사용자 확인 결과, 이번 v3.0에서는 병합하지 않고 각각 독립된 회사로 그대로 처리함. "
        "제조사명 정리(병합) 자체는 이번 범위 밖의 별도 이슈로 남김."
    )

    d.add_paragraph("2. 적용 규칙 (1.41.51 v2.3과 동일, 재설명)").runs[0].bold = True
    d.add_paragraph(
        "공통 수량단위 분리 엔진만 적용됨 — Sigma/Sartorius/Agilent 등 상위20 전용 커스텀 규칙(라벨 "
        "제거, '#'/'[...]' 제거, Agilent UI 전용 처리 등)은 이번 40개 회사명과 일치하는 게 없어 "
        "실질적으로 발동하지 않음. 화이트리스트: KG/MG/UG/NG/AMP/PAK/RXN/TAB/CAP/EA/KT/VL/ML/UL/G/L/%. "
        "복합단위: AMP-EA/KG-K/G-F/SET-F/G-K. NxM 곱셈표기·앞자리 없는 소수·공백하이픈·선두 인치기호도 "
        "동일하게 인식."
    )

    d.add_paragraph("3. 회사별 처리 결과").runs[0].bold = True
    table = d.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["제조사", "대상 건수", "분리텍스트 발생"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
    for mfr in ORDER:
        cells = table.add_row().cells
        cells[0].text = mfr
        cells[1].text = f"{total[mfr]:,}"
        cells[2].text = f"{sep_cnt[mfr]:,}" if sep_cnt[mfr] else "-"
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    d.add_paragraph("")
    d.add_paragraph("4. 검증").runs[0].bold = True
    d.add_paragraph(
        f"대상 건수 합계 {sum(total.values()):,}건, 분리텍스트 발생 총 {sum(sep_cnt.values()):,}건. "
        "분리된 값 전부를 회사별로 육안 검토해서, Agilent·Corning 때처럼 서로 다른 실제 제품을 "
        "잘못 뭉개는 사례가 없음을 확인함(전부 '코드-수량+단위' 형태의 정상적인 분리)."
    )

    d.add_paragraph("")
    d.add_paragraph("5. 컬럼 설명").runs[0].bold = True
    d.add_paragraph("품번_To-Be: 규칙 적용 후 확정된 품번 값. 이번 40개 회사가 아니거나 품번이 없으면 공란.")
    d.add_paragraph("품번_To-Be_비고: 이번 40개 회사 중에는 해당 없음(Sigma 계열 전용 항목이라 발생하지 않음).")
    d.add_paragraph("품번_To-Be_분리텍스트: 원본 품번에서 잘려나간 부분을 그대로 기록.")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    d.save(OUT)
    print(f"[저장] {OUT}")


if __name__ == "__main__":
    main()
