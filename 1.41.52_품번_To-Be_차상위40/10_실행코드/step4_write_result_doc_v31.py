# -*- coding: utf-8 -*-
r"""v3.1 실행 결과 요약 docx 작성. step3의 compute_tobe를 그대로 불러와 실제 데이터로 집계."""
import sys
from pathlib import Path
from collections import Counter

import docx
import pandas as pd
from docx.shared import Pt

sys.path.insert(0, str(Path(__file__).parent))
from step3_build_pn_tobe_v31 import (  # noqa: E402
    SRC, SHEET, HEADER_ROW, COL_PN, COL_MFR_CLEAN_OLD, NEW40_MFRS,
    compute_tobe,
)

OUT = Path(r"C:\Users\정별\1_Work\1.41_동아쏘시오그룹(2)_데이터_표준화\1.41.52_품번_To-Be_차상위40\30_보고서\260812_품번_To-Be_차상위40_결과요약_v3.1.docx")

ORDER = [
    "Charles River", "Km", "R&D Systems", "Combi-Blocks", "Abcam", "PerkinElmer", "대정화금",
    "BioLegend", "SPL", "Saint-Gobain", "Supelco", "싸토리우스코리아", "영인에스티", "TCI", "비티알",
    "DURAN", "Promega", "DAE", "EP", "Cobetter", "MedChemExpress", "ProteinSimple", "GEMÜ",
    "Lonza", "Metrohm", "Avantor", "Beckman Coulter", "Agilent Technologies", "(주)BB", "TRC",
    "Repligen", "Millipore", "Isolab", "Siemens", "Daejung", "Gilson", "DUKSAN", "Rhawn",
    "시너지이노베이션", "Advanced Instruments",
]


def main():
    df = pd.read_excel(SRC, sheet_name=SHEET, header=HEADER_ROW - 1)
    pn_col = df.columns[COL_PN - 1]
    mfr_col = df.columns[COL_MFR_CLEAN_OLD - 1]

    total = Counter()
    sep_cnt = Counter()

    for _, row in df.iterrows():
        pn = row[pn_col]
        mfr = row[mfr_col]
        if pd.isna(pn) or pd.isna(mfr) or mfr not in NEW40_MFRS:
            continue
        total[mfr] += 1
        tobe, memo, sep = compute_tobe(pn, mfr)
        if sep:
            sep_cnt[mfr] += 1

    d = docx.Document()
    t = d.add_paragraph("품번 To-Be — 차상위40(21~60위) v3.1 결과 요약")
    t.runs[0].bold = True
    t.runs[0].font.size = Pt(16)

    d.add_paragraph(
        "산출 파일: 1.41.52_품번_To-Be_차상위40\\20_결과\\"
        "260812_S-TEPS_품번_To-Be_v3.1(전역규칙화_차상위40).xlsx\n"
        "v3.0에서 회사명으로 게이트돼있던 '제거형' 규칙들을 내용 기반 전역 규칙으로 전환해서 재적용."
    )
    d.add_paragraph("")

    d.add_paragraph("1. v3.0 대비 변경 사항").runs[0].bold = True
    d.add_paragraph(
        "① 라벨 프리픽스/서픽스 제거(견적번호/부품 번호/USP/시리얼 번호/모델명/끝 외)를 회사명 조건 "
        "없이, 값이 그 패턴에 맞으면 어느 회사든 적용.\n"
        "② 선두 '#' 제거를 전 회사 공통 적용(기존: 대한과학/Cell Signaling/Sigma/Sigma-Aldrich 전용). "
        "예: DURAN '#21806545'→'21806545', Saint-Gobain '#374-250-3'→'374-250-3'.\n"
        "③ 전체가 '[...]'로 감싸진 경우 제거를 전 회사 공통 적용(기존: Sartorius 전용). 이번 40개 "
        "중에는 해당 건 없음.\n"
        "④ 끝 'UI' 제거를 전 회사 공통 적용(기존: Agilent 전용). 이번 40개 중에는 해당 건 없음.\n"
        "⑤ (전역화하지 않고 그대로 유지) 대한과학·Agilent는 공통 수량단위 분리 엔진 자체를 스킵, "
        "Thermo Fisher는 'KOLAS'로 시작하는 값만 예외 — 이 3개는 특정 회사의 실제 제품코드 구조 때문에 "
        "생긴 안전장치라 사용자 확인 하에 전역화하지 않음.\n"
        "부수적으로 끝 '외' 규칙 적용 중 공백 없이 붙은 케이스(Gilson 'P10외'→'P10')도 함께 발견되어 "
        "반영됨."
    )
    d.add_paragraph(
        "사전 검증: 새 40개 회사 데이터에서 전역화 대상 패턴을 미리 검색한 결과, 실제로 값이 존재한 "
        "것은 선두 '#'(5건: DURAN 3, Saint-Gobain 2)과 끝 '외'(1건, Gilson)뿐이었고 견적번호/부품번호/"
        "USP/시리얼번호/모델명/UI는 0건이라 전역화가 새로운 오탐을 만들지 않음을 확인함."
    )

    d.add_paragraph("2. 처리 범위").runs[0].bold = True
    d.add_paragraph(f"제조사(정리) 21~60위, 총 {sum(total.values()):,}건 (v3.0과 동일 범위).")

    d.add_paragraph("3. 회사별 처리 결과").runs[0].bold = True
    table = d.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["제조사", "대상 건수", "분리텍스트 발생"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
    for mfr in ORDER:
        cells = table.add_row().cells
        cells[0].text = mfr
        cells[1].text = f"{total[mfr]:,}"
        cells[2].text = f"{sep_cnt[mfr]:,}" if sep_cnt[mfr] else "-"
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    d.add_paragraph("")
    d.add_paragraph("4. 검증").runs[0].bold = True
    d.add_paragraph(
        f"대상 건수 합계 {sum(total.values()):,}건(변동 없음), 분리텍스트 발생 총 "
        f"{sum(sep_cnt.values()):,}건(v3.0의 65건 대비 +6건). v3.0과 v3.1의 계산 결과를 전수 대조해서, "
        "이번에 의도한 전역화 항목(#, 끝 외) 외에는 단 한 건도 값이 달라지지 않았음을 확인함 — "
        "총 변경 6건(DURAN 3, Saint-Gobain 2, Gilson 1), 전부 의도한 규칙에 해당."
    )

    d.add_paragraph("")
    d.add_paragraph("5. 컬럼 설명").runs[0].bold = True
    d.add_paragraph("품번_To-Be: 규칙 적용 후 확정된 품번 값. 이번 40개 회사가 아니거나 품번이 없으면 공란.")
    d.add_paragraph("품번_To-Be_비고: 이번 40개 회사 중에는 해당 없음(Sigma 계열 전용 항목).")
    d.add_paragraph("품번_To-Be_분리텍스트: 원본 품번에서 잘려나간 부분을 그대로 기록.")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    d.save(OUT)
    print(f"[저장] {OUT}")


if __name__ == "__main__":
    main()
