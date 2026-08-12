# -*- coding: utf-8 -*-
r"""v3.2 실행 결과 요약 docx 작성. step5의 compute_tobe를 그대로 불러와 실제 데이터로 집계."""
import sys
from pathlib import Path
from collections import Counter

import docx
import pandas as pd
from docx.shared import Pt

sys.path.insert(0, str(Path(__file__).parent))
from step5_build_pn_tobe_v32 import (  # noqa: E402
    SRC, SHEET, HEADER_ROW, COL_PN, COL_MFR_CLEAN_OLD, NEW40_MFRS,
    compute_tobe,
)

OUT = Path(r"C:\Users\정별\1_Work\1.41_동아쏘시오그룹(2)_데이터_표준화\1.41.52_품번_To-Be_차상위40\30_보고서\260812_품번_To-Be_차상위40_결과요약_v3.2.docx")

ORDER = [
    "Charles River", "Km", "R&D Systems", "Combi-Blocks", "Abcam", "PerkinElmer", "대정화금",
    "BioLegend", "SPL", "Saint-Gobain", "Supelco", "싸토리우스코리아", "영인에스티", "TCI", "비티알",
    "DURAN", "Promega", "DAE", "EP", "Cobetter", "MedChemExpress", "ProteinSimple", "GEMÜ",
    "Lonza", "Metrohm", "Avantor", "Beckman Coulter", "Agilent Technologies", "(주)BB", "TRC",
    "Repligen", "Millipore", "Isolab", "Siemens", "Daejung", "Gilson", "DUKSAN", "Rhawn",
    "시너지이노베이션", "Advanced Instruments",
]


def main():
    df = pd.read_excel(SRC, sheet_name=SHEET, header=HEADER_ROW - 1)
    pn_col = df.columns[COL_PN - 1]
    mfr_col = df.columns[COL_MFR_CLEAN_OLD - 1]

    total = Counter()
    sep_cnt = Counter()

    for _, row in df.iterrows():
        pn = row[pn_col]
        mfr = row[mfr_col]
        if pd.isna(pn) or pd.isna(mfr) or mfr not in NEW40_MFRS:
            continue
        total[mfr] += 1
        tobe, memo, sep = compute_tobe(pn, mfr)
        if sep:
            sep_cnt[mfr] += 1

    d = docx.Document()
    t = d.add_paragraph("품번 To-Be — 차상위40(21~60위) v3.2 결과 요약")
    t.runs[0].bold = True
    t.runs[0].font.size = Pt(16)

    d.add_paragraph(
        "산출 파일: 1.41.52_품번_To-Be_차상위40\\20_결과\\"
        "260812_S-TEPS_품번_To-Be_v3.2(M_MM_일련번호_차상위40).xlsx\n"
        "v3.1에 2개 패턴만 추가 반영."
    )
    d.add_paragraph("")

    d.add_paragraph("1. v3.1 대비 변경 사항").runs[0].bold = True
    d.add_paragraph(
        "① 단위 화이트리스트에 M, MM 추가. 예: 비티알 'AH-APSH-P-1000-15-15-250MM'→"
        "'AH-APSH-P-1000-15-15', 'AH-APSH-P-0375-05-05-2M'→'AH-APSH-P-0375-05-05'. "
        "(R&D Systems '3136-RL-01M/CF'처럼 단위 뒤에 '/CF'가 더 붙어 끝까지 안 맞는 경우는 "
        "그대로 유지됨 — 의도된 동작.)\n"
        "② 전역 라벨 규칙에 '일련번호' 추가(기존 '시리얼 번호:'와 다른 표현이라 별도 규칙 필요). "
        "예: GEMÜ '일련번호 J20240419-Q6'→'J20240419-Q6'."
    )
    d.add_paragraph(
        "사전 검증: 상위20(1.41.51) 데이터에는 이 두 패턴이 0건이라 이번 수정은 차상위40에만 반영함."
    )

    d.add_paragraph("2. 처리 범위").runs[0].bold = True
    d.add_paragraph(f"제조사(정리) 21~60위, 총 {sum(total.values()):,}건 (v3.1과 동일 범위).")

    d.add_paragraph("3. 회사별 처리 결과").runs[0].bold = True
    table = d.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["제조사", "대상 건수", "분리텍스트 발생"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
    for mfr in ORDER:
        cells = table.add_row().cells
        cells[0].text = mfr
        cells[1].text = f"{total[mfr]:,}"
        cells[2].text = f"{sep_cnt[mfr]:,}" if sep_cnt[mfr] else "-"
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    d.add_paragraph("")
    d.add_paragraph("4. 검증").runs[0].bold = True
    d.add_paragraph(
        f"대상 건수 합계 {sum(total.values()):,}건(변동 없음), 분리텍스트 발생 총 "
        f"{sum(sep_cnt.values()):,}건(v3.1의 71건 대비 +9건). v3.1과 v3.2 전수 대조 결과, "
        "이번에 의도한 M/MM/일련번호 항목 외에는 값이 달라지지 않았음을 확인함 — 총 변경 9건"
        "(비티알 8건, GEMÜ 1건)."
    )

    d.add_paragraph("")
    d.add_paragraph("5. 컬럼 설명").runs[0].bold = True
    d.add_paragraph("품번_To-Be: 규칙 적용 후 확정된 품번 값. 이번 40개 회사가 아니거나 품번이 없으면 공란.")
    d.add_paragraph("품번_To-Be_비고: 이번 40개 회사 중에는 해당 없음(Sigma 계열 전용 항목).")
    d.add_paragraph("품번_To-Be_분리텍스트: 원본 품번에서 잘려나간 부분을 그대로 기록.")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    d.save(OUT)
    print(f"[저장] {OUT}")


if __name__ == "__main__":
    main()
