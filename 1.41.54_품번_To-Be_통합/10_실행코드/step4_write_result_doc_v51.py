# -*- coding: utf-8 -*-
r"""v5.1 결과 요약 docx 작성. 쉬운 표현 위주로 작성."""
from pathlib import Path
from collections import Counter

import docx
import openpyxl
from docx.shared import Pt

RESULT_XLSX = Path(r"C:\Users\정별\1_Work\1.41_동아쏘시오그룹(2)_데이터_표준화\1.41.54_품번_To-Be_통합\20_결과\260812_S-TEPS_품번_To-Be_v5.1(의견컬럼_LTLC버그수정).xlsx")
OUT = Path(r"C:\Users\정별\1_Work\1.41_동아쏘시오그룹(2)_데이터_표준화\1.41.54_품번_To-Be_통합\30_보고서\260812_품번_To-Be_통합_결과요약_v5.1.docx")

CAT_DESC = {
    "o>유지": "원본 그대로 둔 값(앞뒤 빈칸 정리 말고는 아무것도 안 바꿈)",
    "x>수량단위 제거": "수량이나 단위(예: -100G, -5EA, -10ML 같은 것)를 뒤에서 잘라낸 값",
    "x>안내라벨 제거": "'견적번호', '시리얼 번호:' 같은 안내 문구를 앞이나 뒤에서 지운 값",
    "x>불용기호 제거": "'#', '[ ]' 같은 필요 없는 기호를 지운 값",
    "x>부가정보 제거": "콤마 뒤에 따라붙은 부가 설명이나 부속품 설명 같은 걸 지운 값",
    "x>상세규격 이동": "품번이 아니라 사실상 규격/치수/용량 표기여서 품번_To-Be는 비우고 원래 값은 분리텍스트 쪽으로 옮긴 것",
    "x>의미없음": "품번이라 보기 어려운 서술형 문구라서 품번_To-Be를 비운 것",
}


def main():
    wb = openpyxl.load_workbook(RESULT_XLSX, data_only=False)
    ws = wb["Steps_중복제거_32359"]

    counter = Counter()
    samples = {}
    total_pn = 0
    for row in ws.iter_rows(min_row=5, max_row=ws.max_row):
        pn = row[14].value
        if pn in (None, ""):
            continue
        total_pn += 1
        op = row[17].value
        if op:
            counter[op] += 1
            if op not in samples:
                samples[op] = (pn, row[16].value, row[18].value)

    d = docx.Document()
    t = d.add_paragraph("품번 To-Be — v5.1 (의견 컬럼 정리) 결과 요약")
    t.runs[0].bold = True
    t.runs[0].font.size = Pt(16)

    d.add_paragraph(
        "산출 파일: 1.41.54_품번_To-Be_통합\\20_결과\\"
        "260812_S-TEPS_품번_To-Be_v5.1(의견컬럼_LTLC버그수정).xlsx\n\n"
        "이번에 바뀐 건 딱 2가지입니다.\n"
        "1) 컬럼 이름을 '품번_To-Be_비고'에서 '품번_To-Be_의견'으로 바꾸고, 내용도 예전에 상사님이 "
        "쓰시던 '품번_의견' 컬럼처럼 'o>유지' / 'x>사유' 짧은 형식으로 다시 썼습니다. 예전처럼 문장으로 "
        "길게 설명하던 방식은 그만두고, 아래 표처럼 몇 가지 짧은 사유로 정리했습니다.\n"
        "2) 이전 버전(1.41.53 v4.1) 보고서에서 'GASTEC 4LT, 1LC, 2LC 같은 값은 품번_To-Be를 비운다'고 "
        "적어놨는데, 실제로는 코드에 반영이 안 돼 있던 걸 발견해서 이번에 실제로 반영했습니다(7건)."
    )
    d.add_paragraph("")

    d.add_paragraph("1. '품번_To-Be_의견' 컬럼이란").runs[0].bold = True
    d.add_paragraph(
        "이 행의 품번_To-Be 값을 만들 때 원본을 건드렸는지 안 건드렸는지, 건드렸다면 왜 건드렸는지를 "
        "한눈에 보이게 짧게 적어둔 컬럼입니다. 'o>유지'면 손 안 댄 것이고, 'x>'로 시작하면 뭔가를 "
        "잘라내거나 비웠다는 뜻입니다."
    )

    d.add_paragraph("2. 사유별 건수와 뜻").runs[0].bold = True
    table = d.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["의견 표시", "건수", "무슨 뜻인지"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
    order = [
        "o>유지", "x>수량단위 제거", "x>불용기호 제거", "x>부가정보 제거",
        "x>안내라벨 제거", "x>상세규격 이동", "x>의미없음",
    ]
    for op in order:
        cells = table.add_row().cells
        cells[0].text = op
        cells[1].text = f"{counter.get(op, 0):,}"
        cells[2].text = CAT_DESC.get(op, "")
    combo_total = sum(n for k, n in counter.items() if k not in order)
    if combo_total:
        cells = table.add_row().cells
        cells[0].text = "그 외(2가지 이상 겹침)"
        cells[1].text = f"{combo_total:,}"
        cells[2].text = (
            "한 값에서 2가지 이상 이유로 손댄 경우 (예: 'x>불용기호+수량단위 제거' = "
            "기호도 지우고 수량단위도 같이 지운 경우)"
        )
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    d.add_paragraph("")
    d.add_paragraph("3. 예시").runs[0].bold = True
    example_order = [
        "x>수량단위 제거", "x>안내라벨 제거", "x>불용기호 제거",
        "x>부가정보 제거", "x>상세규격 이동", "x>의미없음",
    ]
    for op in example_order:
        if op in samples:
            pn, tobe, sep = samples[op]
            tobe_disp = tobe if tobe not in (None, "") else "(비움)"
            d.add_paragraph(f"{op} — 예: '{pn}' → '{tobe_disp}'")

    d.add_paragraph("")
    d.add_paragraph("4. 전체 건수").runs[0].bold = True
    d.add_paragraph(
        f"품번이 있는 전체 {total_pn:,}건 중 {counter.get('o>유지', 0):,}건은 원본 그대로, "
        f"나머지 {total_pn - counter.get('o>유지', 0):,}건은 이번 작업에서 뭔가를 지우거나 비웠습니다."
    )

    d.add_paragraph("")
    d.add_paragraph("5. 컬럼 설명").runs[0].bold = True
    d.add_paragraph("품번_To-Be: 정리 끝난 최종 품번 값. 품번이 없으면 공란.")
    d.add_paragraph("품번_To-Be_의견: 위 표처럼 'o>유지' 또는 'x>사유' 형식.")
    d.add_paragraph("품번_To-Be_분리텍스트: 원본에서 실제로 잘려나간(혹은 통째로 빠진) 글자를 그대로 적어둔 것.")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    d.save(OUT)
    print(f"[저장] {OUT}")


if __name__ == "__main__":
    main()
