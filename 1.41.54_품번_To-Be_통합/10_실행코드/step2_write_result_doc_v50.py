# -*- coding: utf-8 -*-
r"""v5.0 통합 결과 요약 docx 작성."""
import sys
from pathlib import Path
from collections import Counter

import docx
import pandas as pd
from docx.shared import Pt

sys.path.insert(0, str(Path(__file__).parent))
from step1_merge_v50 import (  # noqa: E402
    SRC, SHEET, HEADER_ROW, COL_PN, COL_MFR_CLEAN_OLD, compute_tobe_routed,
)

OUT = Path(r"C:\Users\정별\1_Work\1.41_동아쏘시오그룹(2)_데이터_표준화\1.41.54_품번_To-Be_통합\30_보고서\260812_품번_To-Be_통합_결과요약_v5.0.docx")


def main():
    df = pd.read_excel(SRC, sheet_name=SHEET, header=HEADER_ROW - 1)
    pn_col = df.columns[COL_PN - 1]
    mfr_col = df.columns[COL_MFR_CLEAN_OLD - 1]
    df2 = df[[pn_col, mfr_col]].dropna(subset=[pn_col])

    src_counter = Counter()
    n_memo = 0
    n_sep = 0
    for _, row in df2.iterrows():
        tobe, memo, sep, source = compute_tobe_routed(row[pn_col], row[mfr_col])
        src_counter[source] += 1
        if memo:
            n_memo += 1
        if sep:
            n_sep += 1

    total = len(df2)

    d = docx.Document()
    t = d.add_paragraph("품번 To-Be — v5.0 통합 결과 요약")
    t.runs[0].bold = True
    t.runs[0].font.size = Pt(16)

    d.add_paragraph(
        "산출 파일: 1.41.54_품번_To-Be_통합\\20_결과\\"
        "260812_S-TEPS_품번_To-Be_v5.0(통합_상위20_차상위40_나머지).xlsx\n"
        "다음 3개 결과를 회사 범위가 서로 겹치지 않는다는 전제 하에 1개 파일로 통합함:\n"
        "  - 1.41.51_품번_To-Be (상위20, v2.3)\n"
        "  - 1.41.52_품번_To-Be_차상위40 (차상위40, v3.2 — 상위20의 누적 규칙 + 라벨/#/[...]/UI 전역화 "
        "+ M/MM·일련번호)\n"
        "  - 1.41.53_품번_To-Be_나머지전체 (나머지 6,722건, v4.1 — 위 누적 규칙 중 대한과학/Agilent/"
        "Thermo 3개 예외만 제외하고 전역 적용 + 수기검토 보완)\n"
        "통합 방식: 3개 파일을 그대로 붙이지 않고, 원본에서 각 행의 제조사(정리)에 맞는 로직을 다시 "
        "실행해서 값을 채움(지금까지의 작업 방식과 동일하게 재현 가능한 코드로 생성). 제조사(정리) 값은 "
        "1글자라도 다르면 별개 회사로 취급하고 병합하지 않음(정성적 판단 배제)."
    )
    d.add_paragraph("")

    d.add_paragraph("1. 검증").runs[0].bold = True
    d.add_paragraph(
        f"통합 결과 파일의 품번_To-Be/분리텍스트 값을 원본 3개 파일과 행 단위로 전수 대조한 결과, "
        f"검증 대상 14,003건(품번_To-Be가 채워진 행) 전부 정확히 일치함(불일치 0건). 나머지 10건은 "
        "v4.1에서 서술형 값으로 판단해 품번_To-Be를 비운 행이라 이 대조에서는 자연 제외됨(별도 확인 "
        "완료된 항목)."
    )

    d.add_paragraph("2. 처리 범위 및 출처별 건수").runs[0].bold = True
    table = d.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "출처"
    hdr[1].text = "건수"
    for cell in hdr:
        cell.paragraphs[0].runs[0].bold = True
    for source in ["상위20(v2.3)", "차상위40(v3.2)", "나머지전체(v4.1)"]:
        cells = table.add_row().cells
        cells[0].text = source
        cells[1].text = f"{src_counter[source]:,}"
    cells = table.add_row().cells
    cells[0].text = "합계"
    cells[1].text = f"{total:,}"

    d.add_paragraph("")
    d.add_paragraph("3. 전체 요약").runs[0].bold = True
    d.add_paragraph(
        f"품번 보유 전체 {total:,}건에 대해 품번_To-Be가 채워짐(v4.1에서 전체 제거로 비운 10건 포함). "
        f"품번_To-Be_비고 표시 {n_memo:,}건(Sigma/Sigma-Aldrich 패턴 불일치 건). "
        f"품번_To-Be_분리텍스트 채워진 행 {n_sep:,}건."
    )

    d.add_paragraph("")
    d.add_paragraph("4. 컬럼 설명").runs[0].bold = True
    d.add_paragraph("품번_To-Be: 규칙 적용 후 확정된 품번 값. 품번이 없으면 공란.")
    d.add_paragraph("품번_To-Be_비고: Sigma/Sigma-Aldrich 중 수량단위 패턴이 안 맞아 원본을 그대로 유지한 경우, 또는 서술형 값으로 판단되어 전체 제외된 경우에 표시.")
    d.add_paragraph("품번_To-Be_분리텍스트: 원본 품번에서 잘려나간(또는 전체 제외된) 부분을 그대로 기록.")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    d.save(OUT)
    print(f"[저장] {OUT}")


if __name__ == "__main__":
    main()
