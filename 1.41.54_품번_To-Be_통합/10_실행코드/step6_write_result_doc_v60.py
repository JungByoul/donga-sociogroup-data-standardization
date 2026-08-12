# -*- coding: utf-8 -*-
r"""v6.0 결과 요약 docx 작성. 쉬운 표현 위주."""
from pathlib import Path
from collections import Counter

import docx
import openpyxl
from docx.shared import Pt

RESULT_XLSX = Path(r"C:\Users\정별\1_Work\1.41_동아쏘시오그룹(2)_데이터_표준화\1.41.54_품번_To-Be_통합\20_결과\260812_S-TEPS_품번_To-Be_v6.0(4차판정컬럼추가).xlsx")
OUT = Path(r"C:\Users\정별\1_Work\1.41_동아쏘시오그룹(2)_데이터_표준화\1.41.54_품번_To-Be_통합\30_보고서\260812_품번_To-Be_통합_결과요약_v6.0.docx")

REASON_DESC = {
    "x>값없음": "원래 값이 진짜 품번이라 보기엔 의미가 부족하다고 이전 작업(4차판정)에서 판단한 경우",
    "x>설명문키워드": "설명하는 문장/단어라서 품번이 아니라고 이전 작업에서 판단한 경우",
    "x>CAS의심": "화학물질 CAS 번호로 보여서 품번이 아니라고 이전 작업에서 판단한 경우",
    "x>품목명동일": "품목명(품목 이름)과 똑같아서 품번이 아니라고 이전 작업에서 판단한 경우",
    "x>숫자없음": "숫자가 하나도 없어서 품번이 아닐 가능성이 크다고 이전 작업에서 판단한 경우",
    "x>단위값": "'25mg', '500mg'처럼 수량/단위만 있고 품번이 아니라고 이전 작업에서 판단한 경우",
    "x>구조식의심": "화학구조식으로 보여서 품번이 아니라고 이전 작업에서 판단한 경우",
}


def main():
    wb = openpyxl.load_workbook(RESULT_XLSX, data_only=False)
    ws = wb["Steps_중복제거_32359"]

    counter = Counter()
    total_pn = 0
    n_reject = 0
    for row in ws.iter_rows(min_row=5, max_row=ws.max_row):
        pn = row[14].value
        if pn in (None, ""):
            continue
        total_pn += 1
        op = row[17].value
        if op:
            counter[op] += 1
            if op in REASON_DESC:
                n_reject += 1

    d = docx.Document()
    t = d.add_paragraph("품번 To-Be — v6.0 (4차판정 컬럼 추가) 결과 요약")
    t.runs[0].bold = True
    t.runs[0].font.size = Pt(16)

    d.add_paragraph(
        "산출 파일: 1.41.54_품번_To-Be_통합\\20_결과\\"
        "260812_S-TEPS_품번_To-Be_v6.0(4차판정컬럼추가).xlsx\n\n"
        "이번에는 예전에 진행했던 '4차판정' 작업(1.41.44 폴더, 8월 10일 결과물)의 판정 결과를 이 파일에 "
        "가져와 붙였습니다. 그 작업은 품번처럼 보이는 값들 중에서 '진짜 품번이 아닌 것'을 몇 단계에 "
        "걸쳐 걸러낸 작업이었는데, 지금까지 저희가 만든 v1.0~v5.1은 이 판정을 전혀 반영하지 않고 품번이 "
        "있는 값을 전부 정리 대상으로 삼았습니다. 이번 v6.0에서 그 부분을 바로잡았습니다."
    )
    d.add_paragraph("")

    d.add_paragraph("1. 두 파일을 어떻게 연결했는지").runs[0].bold = True
    d.add_paragraph(
        "4차판정 파일과 지금 쓰는 원본 파일은 행 순서가 서로 달라서(정렬이 바뀜), 행 번호로는 연결할 "
        "수 없었습니다. 대신 각 행에 붙어있는 고유번호(A열 'key')로 정확히 짝을 맞춰서 연결했습니다. "
        "품번이 있는 14,012건 전부 짝이 맞는 것을 확인했습니다."
    )

    d.add_paragraph("2. 새로 추가된 컬럼 4개").runs[0].bold = True
    d.add_paragraph(
        "O열_1차판정 / O열_2차판정 / O열_3차판정 / O열_4차판정 — 예전 4차판정 작업에서 나온 판정 결과를 "
        "그대로 옮겨온 것입니다('4차_검토메모' 컬럼은 이번에 가져오지 않았습니다). 이 4개 컬럼 중 하나라도 "
        "'진짜 품번 아님'이라는 표시가 있으면, 그 사유를 아래 3번처럼 처리했습니다."
    )

    d.add_paragraph("3. 이미 '품번 아님'으로 판정된 값은 어떻게 처리했는지").runs[0].bold = True
    d.add_paragraph(
        f"총 {n_reject:,}건이 4차판정에서 이미 '품번이 아니다'라고 판정된 값이었습니다. 이 건들은:\n"
        "- 품번_To-Be 값을 비웠습니다(원래 품번이 아니라고 판정났으니 정리할 필요가 없다고 봄).\n"
        "- 원래 값은 품번_To-Be_분리텍스트 칸으로 그대로 옮겨서 남겨뒀습니다.\n"
        "- 품번_To-Be_의견 칸에는 저희가 만든 사유(수량단위 제거 등) 대신, 4차판정에서 나온 사유를 "
        "짧게 적었습니다."
    )

    d.add_paragraph("4. 4차판정 사유별 건수").runs[0].bold = True
    table = d.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["의견 표시", "건수", "무슨 뜻인지"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
    for op, desc in REASON_DESC.items():
        n = counter.get(op, 0)
        if n == 0:
            continue
        cells = table.add_row().cells
        cells[0].text = op
        cells[1].text = f"{n:,}"
        cells[2].text = desc
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    d.add_paragraph("")
    d.add_paragraph("5. 전체 건수").runs[0].bold = True
    d.add_paragraph(
        f"품번이 있는 전체 {total_pn:,}건 중 {n_reject:,}건은 이번에 4차판정 결과를 따라 품번_To-Be를 "
        f"비웠고, 나머지 {total_pn - n_reject:,}건은 v5.1과 똑같은 방식(o>유지 / x>사유)으로 처리했습니다. "
        "v5.1과 겹치는 부분(4차판정에서 안 걸린 나머지 건들)은 값이 하나도 안 바뀌었다는 것도 전수 대조로 "
        "확인했습니다."
    )

    d.add_paragraph("")
    d.add_paragraph("6. 컬럼 설명").runs[0].bold = True
    d.add_paragraph("품번_To-Be: 정리 끝난 최종 품번 값. 4차판정에서 이미 품번이 아니라고 나온 경우 공란.")
    d.add_paragraph("품번_To-Be_의견: 'o>유지' / 'x>사유' 형식. 4차판정에서 걸린 건은 4차판정 사유로 표시.")
    d.add_paragraph("품번_To-Be_분리텍스트: 잘려나간(또는 통째로 빠진) 글자를 그대로 적어둔 것.")
    d.add_paragraph("O열_1차판정~4차판정: 예전 4차판정 작업 결과를 그대로 가져온 참고용 컬럼.")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    d.save(OUT)
    print(f"[저장] {OUT}")


if __name__ == "__main__":
    main()
