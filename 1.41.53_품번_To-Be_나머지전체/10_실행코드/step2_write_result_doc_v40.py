# -*- coding: utf-8 -*-
r"""v4.0 실행 결과 요약 docx 작성. step1의 compute_tobe를 그대로 불러와 실제 데이터로 집계."""
import sys
from pathlib import Path
from collections import Counter

import docx
import pandas as pd
from docx.shared import Pt

sys.path.insert(0, str(Path(__file__).parent))
from step1_build_pn_tobe_v40 import (  # noqa: E402
    SRC, SHEET, HEADER_ROW, COL_PN, COL_MFR_CLEAN_OLD, TOP60_MFRS,
    compute_tobe,
)

OUT = Path(r"C:\Users\정별\1_Work\1.41_동아쏘시오그룹(2)_데이터_표준화\1.41.53_품번_To-Be_나머지전체\30_보고서\260812_품번_To-Be_나머지전체_결과요약_v4.0.docx")


def main():
    df = pd.read_excel(SRC, sheet_name=SHEET, header=HEADER_ROW - 1)
    pn_col = df.columns[COL_PN - 1]
    mfr_col = df.columns[COL_MFR_CLEAN_OLD - 1]
    df2 = df[[pn_col, mfr_col]].dropna(subset=[pn_col])
    sub = df2[~df2[mfr_col].isin(TOP60_MFRS)]

    total = len(sub)
    sep_cnt = Counter()
    n_sep_total = 0
    for _, row in sub.iterrows():
        pn = row[pn_col]
        mfr = row[mfr_col] if pd.notna(row[mfr_col]) else "(제조사 미기재)"
        tobe, sep = compute_tobe(pn)
        if sep:
            sep_cnt[mfr] += 1
            n_sep_total += 1

    d = docx.Document()
    t = d.add_paragraph("품번 To-Be — 나머지 전체(61위 이하 + 제조사 미기재) v4.0 결과 요약")
    t.runs[0].bold = True
    t.runs[0].font.size = Pt(16)

    d.add_paragraph(
        "산출 파일: 1.41.53_품번_To-Be_나머지전체\\20_결과\\"
        "260812_S-TEPS_품번_To-Be_v4.0(전역규칙_나머지전체).xlsx\n"
        "상위20(1.41.51, v2.3)·차상위40(1.41.52, v3.2)에서 완성된 '전역 규칙'만 그대로 재사용해서, "
        "나머지 전체(제조사(정리) 61위 이하 + 제조사 미기재 포함)에 적용. 이미 완료된 상위60(7,291건)은 "
        "이번 범위에서 완전히 제외하고 손대지 않음."
    )
    d.add_paragraph("")

    d.add_paragraph("1. 적용 규칙").runs[0].bold = True
    d.add_paragraph(
        "라벨 프리픽스/서픽스 제거(견적번호/부품 번호/USP/시리얼 번호/일련번호/모델명/끝 외), 선두 '#' "
        "제거, 전체 '[...]' 감싸짐 제거, 끝 'UI' 제거, 선두 '숫자+인치기호' 제거, 공통 수량단위 분리"
        "(화이트리스트 KG/MG/UG/NG/AMP/PAK/RXN/TAB/CAP/EA/KT/VL/MM/ML/UL/G/L/M/%, 복합단위 "
        "AMP-EA/KG-K/G-F/SET-F/G-K, NxM 곱셈표기, 공백하이픈, 앞자리 없는 소수) — 전부 회사명 무관 "
        "전역 적용."
    )
    d.add_paragraph(
        "제외한 것: 대한과학/Agilent의 공통 수량단위 분리 스킵 예외, Thermo Fisher 'KOLAS' 예외 — "
        "이 3개는 해당 회사에만 있던 특정 위험 사례라 이번 범위(해당 회사 없음)에는 일반화하지 않음."
    )

    d.add_paragraph("2. 발견 및 수정한 오류").runs[0].bold = True
    d.add_paragraph(
        "1차 실행 결과를 전수 스캔한 결과, 'USP' 라벨 규칙(원래 상위20의 'USP'라는 회사명 — 미국약전 "
        "표준물질 — 데이터용으로 만든 규칙, 예: 'USP 1614002')이 공백 없이 'USP'로 시작하는 다른 회사 "
        "'Monucla'의 자체 카탈로그 코드(예: 'USP-110-2107')까지 잘못 잘라서 앞에 하이픈만 남은 깨진 "
        "값('-110-2107')을 만드는 것을 발견함. 정규식을 'USP 다음에 공백이 최소 1개 있어야 함'으로 "
        "고쳐서 재실행 — Monucla 8건이 원본 그대로 복구됨. (상위20/차상위40 결과는 이 문제와 무관함을 "
        "확인해서 그대로 유지, v3.1/v3.2는 수정하지 않음.)"
    )
    d.add_paragraph(
        "그 외 '-' 단독 값(DFE/Operon/케이엠/YOUNG IN ST/제조사 미기재 등 7건)과 URL이 그대로 입력된 "
        "값(Cantabria labs 1건)은 원본 데이터 자체가 그런 형태라 저희 처리 로직과 무관하며, 그대로 "
        "원본 유지됨(정상 동작)."
    )

    d.add_paragraph("3. 처리 범위").runs[0].bold = True
    d.add_paragraph(f"품번이 있고 상위60에 속하지 않는 모든 행, 총 {total:,}건.")

    d.add_paragraph("4. 검증").runs[0].bold = True
    d.add_paragraph(
        f"분리텍스트가 채워진 행은 총 {n_sep_total:,}건, {len(sep_cnt):,}개 회사(제조사 미기재 포함)에 "
        "걸쳐 분포함. 전체 결과를 '앞/뒤에 하이픈이 남거나 빈 값이 되는 등 형태가 깨진 경우'로 자동 "
        "스캔해서 재확인한 결과 0건임을 확인함. 다만 상위60처럼 회사 하나하나를 사람이 직접 검토하지는 "
        "않았으므로(회사 수가 많아 이번 단계에서는 자동 스캔까지만 수행), 사용자가 수기로 검토하며 "
        "추가 보완할 것을 전제로 한 1차 결과임."
    )

    d.add_paragraph("")
    d.add_paragraph("5. 컬럼 설명").runs[0].bold = True
    d.add_paragraph("품번_To-Be: 규칙 적용 후 확정된 품번 값. 상위60이거나 품번이 없으면 공란.")
    d.add_paragraph("품번_To-Be_비고: 이번 범위에는 해당 없음(Sigma 계열 전용 항목, 상위60에만 존재).")
    d.add_paragraph("품번_To-Be_분리텍스트: 원본 품번에서 잘려나간 부분을 그대로 기록.")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    d.save(OUT)
    print(f"[저장] {OUT}")


if __name__ == "__main__":
    main()
