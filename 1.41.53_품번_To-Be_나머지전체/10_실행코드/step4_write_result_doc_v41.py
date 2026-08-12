# -*- coding: utf-8 -*-
r"""v4.1 실행 결과 요약 docx 작성."""
import sys
from pathlib import Path

import docx
import pandas as pd
from docx.shared import Pt

sys.path.insert(0, str(Path(__file__).parent))
from step3_build_pn_tobe_v41 import (  # noqa: E402
    SRC, SHEET, HEADER_ROW, COL_PN, COL_MFR_CLEAN_OLD, TOP60_MFRS,
    INDIVIDUAL_SUFFIX_OVERRIDES, BLANK_ENTIRELY,
    compute_tobe,
)

OUT = Path(r"C:\Users\정별\1_Work\1.41_동아쏘시오그룹(2)_데이터_표준화\1.41.53_품번_To-Be_나머지전체\30_보고서\260812_품번_To-Be_나머지전체_결과요약_v4.1.docx")


def main():
    df = pd.read_excel(SRC, sheet_name=SHEET, header=HEADER_ROW - 1)
    pn_col = df.columns[COL_PN - 1]
    mfr_col = df.columns[COL_MFR_CLEAN_OLD - 1]
    df2 = df[[pn_col, mfr_col]].dropna(subset=[pn_col])
    sub = df2[~df2[mfr_col].isin(TOP60_MFRS)]

    total = len(sub)
    n_sep = 0
    n_blank = 0
    for _, row in sub.iterrows():
        tobe, sep, memo = compute_tobe(row[pn_col])
        if sep:
            n_sep += 1
        if memo:
            n_blank += 1

    d = docx.Document()
    t = d.add_paragraph("품번 To-Be — 나머지 전체 v4.1 결과 요약 (6,722건 수기검토 반영)")
    t.runs[0].bold = True
    t.runs[0].font.size = Pt(16)

    d.add_paragraph(
        "산출 파일: 1.41.53_품번_To-Be_나머지전체\\20_결과\\"
        "260812_S-TEPS_품번_To-Be_v4.1(수기검토반영_나머지전체).xlsx\n"
        "v4.0(전역 규칙 1차 적용) 결과를 사용자가 6,722건 전수 육안 검토한 뒤 전달한 보완 사항을 반영."
    )
    d.add_paragraph("")

    d.add_paragraph("1. 일반 규칙(전역 적용, 7개)").runs[0].bold = True
    d.add_paragraph(
        "① 끝에 콤마+텍스트가 붙은 경우, 마지막 콤마 이후 부분만 1회 제거. 예: '#3662-4000, Nalgene' "
        "→ '3662-4000', 'DH.BOG011,대한과학' → 'DH.BOG011'. (콤마가 여러 번 나오는 값은 마지막 1개만 "
        "제거되므로 일부는 여전히 콤마가 남을 수 있음 — 예: 'BLUE SERIES,5/16X 12 FT,240V' → "
        "'BLUE SERIES,5/16X 12 FT'. 추가 정리가 필요하면 후속 버전에서 반영 예정.)\n"
        "② 끝 '_Unstained' 제거 (TissueArray 14건).\n"
        "③ 복합단위 'ML-R' 추가 (Fluka '34828-40ML-R' → '34828').\n"
        "④ '-숫자중'(한글 단위) 제거 (Komed '500551-1중' → '500551').\n"
        "⑤ 수량단위 접미사 뒤에 '(브랜드명)' 괄호가 더 붙은 경우 함께 분리 "
        "(Avestin 'Z373427-50EA(Merck)' → 'Z373427', 2건).\n"
        "⑥ 전역 라벨 규칙에 'Serial number:' 추가(영문판, 기존 한글 규칙과 별개). "
        "(BioTek Instruments 'Serial number: 1406266' → '1406266')\n"
        "⑦ 값 전체가 '숫자+LT' 또는 '숫자+LC'인 경우 코드 자체가 없는 용기규격 라벨로 판단해 "
        "품번_To-Be를 통째로 비움 (GASTEC '4LT'×3, '1LC'×2, '2LC'×2, 총 7건)."
    )

    d.add_paragraph("2. 개별 처리 항목 (패턴 일반화 위험 있어 실제 발견된 값에만 적용)").runs[0].bold = True
    p = d.add_paragraph()
    p.add_run(f"접미 분리 {len(INDIVIDUAL_SUFFIX_OVERRIDES)}건: ").bold = True
    p.add_run(", ".join(f"'{k}'→'{v[0]}'" for k, v in INDIVIDUAL_SUFFIX_OVERRIDES.items()))
    p2 = d.add_paragraph()
    p2.add_run(f"전체 제거(품번_To-Be 통째로 비움) {len(BLANK_ENTIRELY)}건: ").bold = True
    p2.add_run(", ".join(f"'{v}'" for v in BLANK_ENTIRELY))
    d.add_paragraph(
        "※ 이 항목들은 정확히 이 원본값에 대해서만 적용되는 화이트리스트 방식이며, 비슷하게 생긴 다른 "
        "값에는 일반화되지 않음. 예: '-한글' 접미사 제거는 nilfisk 3건만 적용했고, 같은 모양(하이픈+한글)인 "
        "'경인에스브이씨 UN-나인나인', 'UN-쿨톤'은 색상/모델명일 가능성이 있어 그대로 유지."
    )

    d.add_paragraph("3. 처리 범위").runs[0].bold = True
    d.add_paragraph(f"품번이 있고 상위60에 속하지 않는 모든 행, 총 {total:,}건 (v4.0과 동일 범위).")

    d.add_paragraph("4. 검증").runs[0].bold = True
    d.add_paragraph(
        f"분리텍스트 발생 총 {n_sep:,}건(v4.0의 247건 대비 +61건), 전체 제거(비움) {n_blank:,}건. "
        "v4.0과 v4.1을 전수 대조해서 총 62건이 바뀌었고 전부 이번에 의도한 규칙(위 1·2번 항목)에만 "
        "해당함을 확인함. 결과가 깨진 형태(하이픈만 남거나 등)로 새로 생긴 경우는 없음."
    )

    d.add_paragraph("")
    d.add_paragraph("5. 컬럼 설명").runs[0].bold = True
    d.add_paragraph(
        "품번_To-Be: 규칙 적용 후 확정된 품번 값. 상위60이거나 품번이 없으면 공란. 서술형 값/규격으로 "
        "판단된 10건은 이번에 통째로 비움(품번_To-Be_비고에 사유 표시)."
    )
    d.add_paragraph("품번_To-Be_비고: 전체 제거된 10건에 한해 '서술형 값/규격으로 판단되어 제외' 표시.")
    d.add_paragraph("품번_To-Be_분리텍스트: 원본 품번에서 잘려나간(또는 전체 제거된) 부분을 그대로 기록.")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    d.save(OUT)
    print(f"[저장] {OUT}")


if __name__ == "__main__":
    main()
