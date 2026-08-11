# -*- coding: utf-8 -*-
r"""
라벨 설명서 v1.0 -> v2.0

v2.3 반영: [코드뭉침]/[설명문혼입] 라벨과 슬래시포함(수동검토필요) 검토메모를
패턴 기반 자동판정이라는 이유로 전부 폐기했으므로, 설명서에서도 해당 항목을 제거하고
4차판정에 [단위값]/[구조식의심](둘 다 육안확인 하드코딩)만 남았음을 반영한다.
"""
from pathlib import Path

import docx

SRC = Path(r"C:\Users\정별\1_Work\1.41_동아쏘시오그룹(2)_데이터_표준화\1.41.44_품번_4차분리\30_보고서\260810_O열_라벨_설명서_v1.0.docx")
DST = Path(r"C:\Users\정별\1_Work\1.41_동아쏘시오그룹(2)_데이터_표준화\1.41.44_품번_4차분리\30_보고서\260810_O열_라벨_설명서_v2.0.docx")


def remove_row(table, idx):
    row = table.rows[idx]
    row._tr.getparent().remove(row._tr)


def main():
    d = docx.Document(SRC)

    d.paragraphs[1].text = "1.41.44_품번_4차분리 · v2.3 기준 — 라벨을 볼 때마다 이 문서만 확인하면 됨"

    # 검토메모(슬래시) 관련 안내 문단은 더 이상 유효하지 않으므로 제거
    p15 = d.paragraphs[15]
    p15._p.getparent().remove(p15._p)

    table3 = d.tables[3]
    # 뒤에서부터 제거해야 인덱스가 안 밀림: row5(슬래시메모) -> row2([설명문혼입]) -> row1([코드뭉침])
    remove_row(table3, 5)
    remove_row(table3, 2)
    remove_row(table3, 1)

    DST.parent.mkdir(parents=True, exist_ok=True)
    d.save(DST)
    print(f"[저장] {DST}")


if __name__ == "__main__":
    main()
